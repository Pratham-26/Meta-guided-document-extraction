from pathlib import Path

import pytest

from src.utils.text import extract_text_from_file


def _create_docx(path: Path, text: str):
    from docx import Document

    doc = Document()
    doc.add_paragraph(text)
    doc.save(str(path))


def _create_odt(path: Path, text: str):
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    doc = OpenDocumentText()
    p = P(text=text)
    doc.text.addElement(p)
    doc.save(str(path))


class TestExtractTextFromPlainText:
    def test_txt_file(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("Hello world", encoding="utf-8")
        assert extract_text_from_file(f) == "Hello world"

    def test_md_file(self, tmp_path):
        f = tmp_path / "test.md"
        f.write_text("# Heading\nSome text", encoding="utf-8")
        result = extract_text_from_file(f)
        assert "Heading" in result
        assert "Some text" in result

    def test_csv_file(self, tmp_path):
        f = tmp_path / "test.csv"
        f.write_text("a,b,c\n1,2,3", encoding="utf-8")
        result = extract_text_from_file(f)
        assert "a,b,c" in result

    def test_unknown_extension_falls_back_to_plain_text(self, tmp_path):
        f = tmp_path / "test.custom"
        f.write_text("custom format content", encoding="utf-8")
        assert extract_text_from_file(f) == "custom format content"


class TestExtractTextFromDocx:
    def test_docx_file(self, tmp_path):
        f = tmp_path / "test.docx"
        _create_docx(f, "This is a Word document")
        result = extract_text_from_file(f)
        assert "This is a Word document" in result

    def test_docx_multiple_paragraphs(self, tmp_path):
        from docx import Document

        f = tmp_path / "multi.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        doc.add_paragraph("")
        doc.save(str(f))
        result = extract_text_from_file(f)
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_corrupted_docx_raises(self, tmp_path):
        f = tmp_path / "bad.docx"
        f.write_bytes(b"not a real docx file content here")
        with pytest.raises(ValueError, match="Failed to extract text"):
            extract_text_from_file(f)


class TestExtractTextFromHtml:
    def test_html_file(self, tmp_path):
        f = tmp_path / "test.html"
        f.write_text(
            "<html><body><h1>Title</h1><p>Content here</p></body></html>",
            encoding="utf-8",
        )
        result = extract_text_from_file(f)
        assert "Title" in result
        assert "Content here" in result

    def test_htm_extension(self, tmp_path):
        f = tmp_path / "test.htm"
        f.write_text(
            "<html><body><p>HTM file content</p></body></html>",
            encoding="utf-8",
        )
        result = extract_text_from_file(f)
        assert "HTM file content" in result


class TestExtractTextFromRtf:
    def test_rtf_file(self, tmp_path):
        f = tmp_path / "test.rtf"
        f.write_text(
            r"{\rtf1\ansi{Sample RTF text}}",
            encoding="utf-8",
        )
        result = extract_text_from_file(f)
        assert "Sample RTF text" in result


class TestExtractTextFromOdt:
    def test_odt_file(self, tmp_path):
        f = tmp_path / "test.odt"
        _create_odt(f, "OpenDocument text content")
        result = extract_text_from_file(f)
        assert "OpenDocument text content" in result


class TestExtractTextErrors:
    def test_missing_file_raises(self, tmp_path):
        f = tmp_path / "nonexistent.txt"
        with pytest.raises(ValueError, match="File not found"):
            extract_text_from_file(f)

    def test_binary_file_raises(self, tmp_path):
        f = tmp_path / "test.xyz"
        f.write_bytes(bytes(range(256)))
        with pytest.raises(ValueError, match="Cannot read"):
            extract_text_from_file(f)
